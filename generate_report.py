from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Set margins
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(2.54)

def add_heading_custom(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_centered_text(text, bold=False, size=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p

def add_normal_para(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    return p

SCREENSHOT_DIR = '/projects/sandbox/summa2'

def add_figure_placeholder(caption, image_file=None):
    """If image_file is provided and exists, embed the actual screenshot.
    Otherwise leave a text placeholder for the user to insert one later."""
    if image_file:
        img_path = os.path.join(SCREENSHOT_DIR, image_file)
        if os.path.exists(img_path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_path, width=Inches(5.5))
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            crun = cap.add_run(caption)
            crun.italic = True
            crun.font.size = Pt(11)
            return cap
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'\n{caption}\n(Insert Screenshot Here)\n')
    run.font.size = Pt(11)
    run.italic = True
    return p


# ============ COVER PAGE ============
add_centered_text('\n\n\n', size=12)
add_centered_text('APJ ABDUL KALAM TECHNOLOGICAL UNIVERSITY', bold=True, size=16)
add_centered_text('(Affiliated Institution)', size=12)
add_centered_text('\n', size=12)
add_centered_text('DEPARTMENT OF COMPUTER SCIENCE AND ENGINEERING', bold=True, size=13)
add_centered_text('\n', size=12)
add_centered_text('MINI PROJECT REPORT', bold=True, size=16)
add_centered_text('On', size=12)
add_centered_text('\n', size=8)
add_centered_text('FinGenius AI - Smart Personal Finance Tracker', bold=True, size=16)
add_centered_text('\n', size=12)
add_centered_text('Submitted in partial fulfillment of the requirements for the degree of', size=12)
add_centered_text('Master of Technology', bold=True, size=13)
add_centered_text('in Computer Science and Engineering', size=12)
add_centered_text('\n', size=12)
add_centered_text('Submitted by:', size=12)
add_centered_text('[Student Name]', bold=True, size=13)
add_centered_text('Register No: [Register Number]', size=12)
add_centered_text('\n', size=12)
add_centered_text('Under the Guidance of', size=12)
add_centered_text('[Guide Name]', bold=True, size=13)
add_centered_text('[Designation]', size=12)
add_centered_text('\n\n', size=12)
add_centered_text('2025-2026', bold=True, size=14)
doc.add_page_break()


# ============ DECLARATION ============
add_heading_custom('DECLARATION', level=1)
add_normal_para('')
add_normal_para(
    'I hereby declare that this Mini Project Report entitled "FinGenius AI - Smart Personal '
    'Finance Tracker" is a bonafide record of the work done by me during the academic year '
    '2025-2026 under the guidance of [Guide Name], [Designation], Department of Computer '
    'Science and Engineering, [College Name].'
)
add_normal_para(
    'I further declare that this project work has not been submitted to any other university '
    'or institution for the award of any degree or diploma.'
)
add_normal_para('\n\n')
add_normal_para('Place: [City]')
add_normal_para('Date: [Date]')
add_normal_para('\n\n')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('[Student Name]\nRegister No: [Register Number]')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
doc.add_page_break()


# ============ ACKNOWLEDGEMENT ============
add_heading_custom('ACKNOWLEDGEMENT', level=1)
add_normal_para('')
add_normal_para(
    'I would like to express my sincere gratitude to all those who helped me in completing '
    'this mini project successfully.'
)
add_normal_para(
    'First and foremost, I thank God Almighty for giving me the strength and wisdom to '
    'complete this work on time.'
)
add_normal_para(
    'I am deeply grateful to my project guide, [Guide Name], [Designation], Department of '
    'Computer Science and Engineering, for the constant encouragement, valuable suggestions, '
    'and guidance throughout the development of this project.'
)
add_normal_para(
    'I express my heartfelt thanks to [HOD Name], Head of the Department of Computer Science '
    'and Engineering, for providing the necessary facilities and support.'
)
add_normal_para(
    'I am also thankful to [Principal Name], Principal of [College Name], for the '
    'encouragement and institutional support.'
)
add_normal_para(
    'Finally, I thank my family and friends for their constant support and motivation '
    'throughout this journey.'
)
add_normal_para('\n\n')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('[Student Name]')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
doc.add_page_break()


# ============ ABSTRACT ============
add_heading_custom('ABSTRACT', level=1)
add_normal_para('')
add_normal_para(
    'Managing personal finances is one of the biggest challenges faced by individuals today. '
    'Most people struggle to keep track of their daily expenses, monthly income, and savings '
    'goals. Traditional methods like maintaining a notebook or using basic spreadsheets are '
    'time-consuming and do not provide any intelligent insights about spending patterns.'
)
add_normal_para(
    'This project presents FinGenius AI, a smart personal finance tracker that combines modern '
    'web technologies with artificial intelligence to help users manage their money effectively. '
    'The application is built using the MERN stack (MongoDB, Express.js, React, Node.js) and '
    'integrates Google Gemini AI to provide personalized financial advice based on the user\'s '
    'actual spending data.'
)
add_normal_para(
    'FinGenius AI allows users to track their expenses and income, create category-wise monthly '
    'budgets, set financial goals with target amounts, and visualize their financial health '
    'through interactive charts and analytics. The application also features Receipt OCR scanning '
    'using Tesseract.js to automatically extract expense details from bill images, voice-based '
    'expense entry using the Web Speech API, and an AI-powered chatbot that can answer questions '
    'about the user\'s finances.'
)
add_normal_para(
    'The system uses JWT-based authentication for secure access, MongoDB Atlas for cloud-based '
    'data storage, and implements a responsive glassmorphism UI design with dark and light mode '
    'support. The application is fully responsive and works seamlessly on desktop, tablet, and '
    'mobile devices.'
)
add_normal_para(
    'Key features include real-time dashboard with financial overview, expense and income '
    'management with search and filtering, budget planning with category-wise allocation, '
    'goal tracking with contribution management, six types of analytical charts, AI chatbot '
    'for financial advice, receipt scanning, voice entry, notifications, and data export.'
)
add_normal_para('\n')
add_normal_para('Keywords: Personal Finance, MERN Stack, Artificial Intelligence, Gemini AI, '
    'Receipt OCR, Budget Planning, Expense Tracker, Web Application')
doc.add_page_break()


# ============ TABLE OF CONTENTS ============
add_heading_custom('TABLE OF CONTENTS', level=1)
add_normal_para('')
toc_items = [
    ('', 'Declaration', 'ii'),
    ('', 'Acknowledgement', 'iii'),
    ('', 'Abstract', 'iv'),
    ('', 'Table of Contents', 'v'),
    ('', 'List of Figures', 'vi'),
    ('', 'List of Tables', 'vii'),
    ('1', 'INTRODUCTION', '1'),
    ('1.1', 'Background', '1'),
    ('1.2', 'Problem Statement', '2'),
    ('1.3', 'Objectives', '2'),
    ('1.4', 'Scope', '3'),
    ('2', 'EXISTING SYSTEM', '4'),
    ('2.1', 'Existing Methods', '4'),
    ('2.2', 'Problems in Existing Systems', '5'),
    ('2.3', 'Need for Proposed System', '5'),
    ('3', 'PROPOSED SYSTEM', '6'),
    ('3.1', 'Overview', '6'),
    ('3.2', 'System Architecture', '7'),
    ('3.3', 'Workflow', '8'),
    ('3.4', 'Module Description', '9'),
    ('3.5', 'Database Design', '15'),
    ('3.6', 'Technology Stack', '17'),
    ('4', 'MERITS AND DEMERITS', '18'),
    ('5', 'COMPARISON WITH EXISTING SYSTEM', '19'),
    ('6', 'CONCLUSION AND FUTURE SCOPE', '20'),
    ('', 'References', '21'),
]
for num, title, page in toc_items:
    p = doc.add_paragraph()
    if num:
        p.add_run(f'{num}    {title}').font.name = 'Times New Roman'
    else:
        p.add_run(f'       {title}').font.name = 'Times New Roman'
    p.add_run(f'    ......    {page}').font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(2)
doc.add_page_break()


# ============ CHAPTER 1 - INTRODUCTION ============
add_heading_custom('CHAPTER 1', level=1)
add_heading_custom('INTRODUCTION', level=2)

add_heading_custom('1.1 Background', level=2)
add_normal_para(
    'In today\'s fast-paced world, financial literacy and money management have become essential '
    'life skills. With the rise of digital payments, online shopping, and multiple income sources, '
    'people are finding it increasingly difficult to keep track of where their money goes each '
    'month. According to various surveys, a significant percentage of working professionals in '
    'India do not maintain any record of their monthly expenses, leading to overspending, '
    'insufficient savings, and financial stress.'
)
add_normal_para(
    'The traditional approach of maintaining expense diaries or using basic Excel spreadsheets '
    'requires manual effort and offers no intelligence or insights. While there are several '
    'finance tracking apps available in the market, most of them either require paid subscriptions '
    'for advanced features, have complicated user interfaces, or do not provide AI-based '
    'personalized recommendations. There is a clear gap for an intelligent, user-friendly, and '
    'free personal finance management tool that can not only track transactions but also provide '
    'smart suggestions to improve financial health.'
)
add_normal_para(
    'The rapid advancement of artificial intelligence, particularly large language models like '
    'Google Gemini, has opened new possibilities for building smart applications that can '
    'understand context and provide meaningful advice. By combining modern web development '
    'frameworks like React and Node.js with AI capabilities, it is now possible to create '
    'applications that serve as personal financial advisors accessible to everyone.'
)


add_heading_custom('1.2 Problem Statement', level=2)
add_normal_para(
    'Most individuals, especially students and young professionals, face the following problems '
    'in managing their personal finances:'
)
add_normal_para(
    '- They do not have a centralized platform to record all their income and expenses in one place.'
)
add_normal_para(
    '- Manual entry of expenses is tedious, leading to incomplete records.'
)
add_normal_para(
    '- There is no easy way to scan receipts and automatically extract expense details.'
)
add_normal_para(
    '- Existing free tools lack AI-powered insights and personalized financial advice.'
)
add_normal_para(
    '- Budget tracking and goal setting features are often locked behind paywalls.'
)
add_normal_para(
    '- Users cannot ask natural language questions about their spending patterns.'
)
add_normal_para(
    'There is a need for a comprehensive, intelligent, and free personal finance management '
    'system that addresses all these challenges while being easy to use and accessible across devices.'
)

add_heading_custom('1.3 Objectives', level=2)
add_normal_para('The main objectives of this project are:')
add_normal_para(
    '1. To develop a full-stack web application for personal finance management using the MERN stack.'
)
add_normal_para(
    '2. To implement secure user authentication with JWT tokens for data privacy.'
)
add_normal_para(
    '3. To enable users to track their expenses and income with detailed categorization.'
)
add_normal_para(
    '4. To provide budget planning with category-wise monthly limits and alerts.'
)
add_normal_para(
    '5. To allow users to set financial goals and track their progress over time.'
)
add_normal_para(
    '6. To integrate Google Gemini AI for providing personalized financial advice through a chatbot.'
)
add_normal_para(
    '7. To implement Receipt OCR scanning for automatic expense extraction from bill images.'
)
add_normal_para(
    '8. To provide voice-based expense entry for hands-free operation.'
)
add_normal_para(
    '9. To present financial data through interactive charts and analytics dashboards.'
)
add_normal_para(
    '10. To build a responsive, modern UI that works on all screen sizes.'
)


add_heading_custom('1.4 Scope', level=2)
add_normal_para(
    'The scope of this project covers the complete design and development of a personal finance '
    'tracking web application with the following capabilities:'
)
add_normal_para(
    '- User registration and login with secure authentication.'
)
add_normal_para(
    '- Complete expense and income management with CRUD operations.'
)
add_normal_para(
    '- Category-wise budget creation and monitoring.'
)
add_normal_para(
    '- Financial goal setting with contribution tracking.'
)
add_normal_para(
    '- AI-powered chatbot for financial queries and advice.'
)
add_normal_para(
    '- Receipt scanning using OCR technology.'
)
add_normal_para(
    '- Voice-based expense entry using browser speech recognition.'
)
add_normal_para(
    '- Analytics dashboard with multiple chart types.'
)
add_normal_para(
    '- Notification system for budget alerts and goal achievements.'
)
add_normal_para(
    '- Dark mode and light mode UI themes.'
)
add_normal_para(
    '- Responsive design for mobile, tablet, and desktop.'
)
add_normal_para(
    'The application is designed for individual users and does not include multi-user '
    'collaboration, bank account integration, or payment gateway features. It focuses on '
    'manual and semi-automated expense tracking with AI assistance.'
)
doc.add_page_break()


# ============ CHAPTER 2 - EXISTING SYSTEM ============
add_heading_custom('CHAPTER 2', level=1)
add_heading_custom('EXISTING SYSTEM', level=2)

add_heading_custom('2.1 Existing Methods', level=2)
add_normal_para(
    'Currently, individuals use various methods to manage their personal finances. Some of the '
    'commonly used approaches are discussed below.'
)
add_normal_para(
    'Manual Methods: Many people still use pen-and-paper diaries, notebooks, or simple '
    'calculators to track their expenses. This method is basic and does not offer any analysis '
    'or insight capabilities. It is also prone to errors and easy to forget.'
)
add_normal_para(
    'Spreadsheet-Based Tracking: Some users rely on Microsoft Excel or Google Sheets to '
    'maintain their financial records. While this provides better organization than pen and paper, '
    'it requires manual data entry, formula creation, and chart generation. It does not provide '
    'real-time alerts or AI-based suggestions.'
)
add_normal_para(
    'Mobile Banking Apps: Banks provide their own mobile apps that show transaction history. '
    'However, these apps only show transactions from one bank account and do not allow '
    'categorization, budget setting, or goal planning. They also do not track cash expenses.'
)
add_normal_para(
    'Third-Party Finance Apps: Applications like Money Manager, Walnut, and Mint offer '
    'expense tracking features. However, many of these apps require premium subscriptions for '
    'advanced features like budget alerts, analytics, or AI insights. Some also have privacy '
    'concerns as they require access to bank SMS messages.'
)

add_heading_custom('2.2 Problems in Existing Systems', level=2)
add_normal_para(
    'After studying the existing methods, the following problems were identified:'
)
add_normal_para(
    '1. Limited Intelligence: Most free tools simply record transactions without providing '
    'any smart analysis or advice on how to improve spending habits.'
)
add_normal_para(
    '2. No AI Chatbot: Users cannot ask natural language questions like "How much did I spend '
    'on food this month?" or "Can I afford a new laptop?" None of the free tools offer this feature.'
)
add_normal_para(
    '3. No Receipt Scanning: Free versions of most apps do not support scanning receipts to '
    'automatically extract expense details, making manual entry the only option.'
)
add_normal_para(
    '4. No Voice Entry: Very few applications support voice-based transaction entry, which is '
    'a convenient feature for quick logging.'
)
add_normal_para(
    '5. Privacy Concerns: Many apps require access to bank SMS or account details, which '
    'raises security and privacy concerns.'
)
add_normal_para(
    '6. Paid Features: Most useful features like analytics, budget alerts, and goal tracking '
    'are locked behind monthly subscriptions.'
)
add_normal_para(
    '7. Poor UI/UX: Many free finance apps have cluttered or outdated interfaces that do not '
    'appeal to younger users.'
)


add_heading_custom('2.3 Need for Proposed System', level=2)
add_normal_para(
    'Considering the limitations of existing systems, there is a strong need for a modern, '
    'AI-powered personal finance management tool that is:'
)
add_normal_para(
    '- Completely free with all features accessible without any subscription.'
)
add_normal_para(
    '- Intelligent enough to provide personalized advice based on actual user data.'
)
add_normal_para(
    '- Equipped with modern input methods like receipt scanning and voice entry.'
)
add_normal_para(
    '- Visually appealing with a modern UI design that attracts young users.'
)
add_normal_para(
    '- Secure and privacy-focused, not requiring access to bank accounts or SMS.'
)
add_normal_para(
    '- Accessible from any device through a responsive web interface.'
)
add_normal_para(
    '- Built with scalable modern technologies for future enhancement.'
)
add_normal_para(
    'The proposed system, FinGenius AI, addresses all these needs by combining the MERN stack '
    'with Google Gemini AI, Tesseract.js OCR, and Web Speech API to create a comprehensive '
    'personal finance management solution.'
)
doc.add_page_break()


# ============ CHAPTER 3 - PROPOSED SYSTEM ============
add_heading_custom('CHAPTER 3', level=1)
add_heading_custom('PROPOSED SYSTEM', level=2)

add_heading_custom('3.1 Overview', level=2)
add_normal_para(
    'FinGenius AI is a smart personal finance tracker designed to help users manage their '
    'money efficiently with the help of artificial intelligence. The application provides a '
    'complete suite of financial management tools including expense tracking, income management, '
    'budget planning, goal setting, analytics, AI chat, receipt OCR, and voice entry.'
)
add_normal_para(
    'The system follows a client-server architecture where the React-based frontend communicates '
    'with the Node.js/Express backend through RESTful APIs. All user data is stored securely in '
    'MongoDB Atlas cloud database. The Google Gemini AI model is integrated on the backend to '
    'process user queries along with their financial context and provide meaningful responses.'
)
add_normal_para(
    'The application features a modern glassmorphism UI design with smooth animations, dark and '
    'light theme modes, and a fully responsive layout that adapts to any screen size. Users can '
    'register, log in, and immediately start tracking their finances with an intuitive dashboard '
    'that shows real-time financial summaries, quick action buttons, and charts.'
)

add_heading_custom('3.2 System Architecture', level=2)
add_normal_para(
    'The system follows a three-tier architecture consisting of the Presentation Layer (Frontend), '
    'Business Logic Layer (Backend), and Data Layer (Database). Below is the architectural '
    'diagram of the system:'
)
add_normal_para('')
# Text-based architecture diagram
arch_text = '''
+------------------------------------------------------------------+
|                     PRESENTATION LAYER                             |
|                    (React + Tailwind CSS)                          |
|                                                                    |
|  Landing | Login | Register | Dashboard | Expenses | Income       |
|  Budget | Goals | Analytics | AI Chat | Settings                  |
+------------------------------------------------------------------+
                              |
                         Axios HTTP
                         (REST API)
                              |
+------------------------------------------------------------------+
|                    BUSINESS LOGIC LAYER                            |
|                  (Node.js + Express.js)                            |
|                                                                    |
|  Auth | Expenses | Income | Budget | Goals | Analytics            |
|  Chat | OCR | Profile | Notifications | Dashboard                 |
|                                                                    |
|  Middleware: JWT Auth | Rate Limiter | Validator | Error Handler  |
+------------------------------------------------------------------+
                              |
              +---------------+----------------+
              |               |                |
+-------------+  +------------+--+  +----------+--------+
|  MongoDB     |  |  Gemini AI   |  |  Tesseract.js    |
|  Atlas       |  |  (Google)    |  |  (OCR Engine)    |
|  (Database)  |  |              |  |                   |
+--------------+  +--------------+  +-------------------+
'''
p = doc.add_paragraph()
run = p.add_run(arch_text)
run.font.name = 'Courier New'
run.font.size = Pt(8)

add_figure_placeholder('Figure 3.1 System Architecture Diagram')


add_heading_custom('3.3 Workflow', level=2)
add_normal_para(
    'The workflow of FinGenius AI describes how a user interacts with the system from registration '
    'to daily usage. The following steps explain the complete user journey:'
)
add_normal_para(
    'Step 1 - Registration: A new user visits the landing page and clicks "Get Started". They '
    'fill in their full name, email address, and password to create an account. The password is '
    'hashed using bcrypt before storage.'
)
add_normal_para(
    'Step 2 - Login: The user logs in with their email and password. The server validates the '
    'credentials and returns a JWT token that is stored in the browser for subsequent requests.'
)
add_normal_para(
    'Step 3 - Dashboard: After login, the user is taken to the dashboard which shows their '
    'total balance, monthly income, monthly expenses, and monthly savings with percentage '
    'comparisons to the previous month. Quick action buttons allow instant navigation to key features.'
)
add_normal_para(
    'Step 4 - Adding Transactions: Users can add expenses or income manually through forms, '
    'scan a receipt image for automatic extraction, or use voice input to speak their transaction.'
)
add_normal_para(
    'Step 5 - Budget Planning: Users create monthly budgets by setting spending limits for '
    'different categories. The system tracks actual spending against these limits.'
)
add_normal_para(
    'Step 6 - Goal Setting: Users create financial goals with target amounts and deadlines. '
    'They can make contributions towards goals and track progress.'
)
add_normal_para(
    'Step 7 - Analytics: Users can view their financial data through various charts including '
    'monthly expense trends, income vs expense comparisons, category-wise breakdowns, and '
    'savings trends over 12 months.'
)
add_normal_para(
    'Step 8 - AI Chat: Users can ask the AI chatbot any question about their finances. The '
    'chatbot has access to the user\'s financial summary and provides personalized answers.'
)
add_normal_para('')
# Workflow text diagram
workflow_text = '''
User Registration --> Login --> Dashboard
                                   |
            +----------+-----------+-----------+-----------+
            |          |           |           |           |
        Expenses    Income     Budget       Goals     Analytics
            |          |           |           |           |
            +----------+-----------+-----------+-----------+
                                   |
                              AI Chatbot
                        (Powered by Gemini AI)
'''
p = doc.add_paragraph()
run = p.add_run(workflow_text)
run.font.name = 'Courier New'
run.font.size = Pt(9)

add_figure_placeholder('Figure 3.2 System Workflow Diagram')
doc.add_page_break()


add_heading_custom('3.4 Module Description', level=2)
add_normal_para(
    'This section provides a detailed description of each module implemented in the FinGenius AI '
    'application. Each module is explained with its purpose, functionality, and key features.'
)

# Module 1: Authentication
add_heading_custom('3.4.1 Authentication Module', level=3)
add_normal_para(
    'The Authentication module handles user registration, login, password management, and session '
    'control. It is the gateway to the application and ensures that only authorized users can '
    'access their financial data.'
)
add_normal_para(
    'Registration: New users provide their full name, email address, and password. The system '
    'validates that the email is unique and the password meets minimum length requirements (6 '
    'characters). The password is hashed using bcryptjs with 12 salt rounds before being stored '
    'in the database. This ensures that even if the database is compromised, passwords remain safe.'
)
add_normal_para(
    'Login: Users enter their email and password. The server compares the provided password '
    'with the stored hash. On successful authentication, a JSON Web Token (JWT) is generated '
    'with a 7-day expiry and sent to the client. This token is included in all subsequent API '
    'requests for authentication.'
)
add_normal_para(
    'Forgot Password: Users who forget their password can request a reset link by entering '
    'their email. The system generates a reset token and sends it via email using Nodemailer. '
    'The token has a limited validity period for security.'
)
add_normal_para(
    'Reset Password: Using the token received via email, users can set a new password. The '
    'old password is replaced with the new hashed password in the database.'
)
add_normal_para(
    'Session Management: The JWT token is stored in the browser\'s local storage. An Auth '
    'Context (React Context API) manages the user state across the application. Protected '
    'routes redirect unauthenticated users to the login page.'
)
add_figure_placeholder('Figure 3.3 Landing Page', 'Screenshot 2026-07-13 195649.png')
add_figure_placeholder('Figure 3.4 Registration Page', 'Screenshot 2026-07-13 195756.png')


# Module 2: Dashboard
add_heading_custom('3.4.2 Dashboard Module', level=3)
add_normal_para(
    'The Dashboard is the main landing page after login and provides a comprehensive overview '
    'of the user\'s financial health at a glance. It serves as the central hub from which users '
    'can navigate to any feature of the application.'
)
add_normal_para(
    'Financial Summary Cards: The dashboard displays four summary cards at the top showing '
    'Total Balance, Monthly Income, Monthly Expense, and Monthly Savings. Each card shows the '
    'current value in Indian Rupees and the percentage change compared to the previous month, '
    'making it easy to understand financial trends.'
)
add_normal_para(
    'Quick Actions: Below the summary cards, four quick action buttons are provided for the '
    'most common tasks - Add Expense, Add Income, AI Chat, and Scan Receipt. These buttons '
    'allow users to perform key operations without navigating to separate pages.'
)
add_normal_para(
    'Charts: The dashboard includes an Income vs Expense line chart and an Expense by Category '
    'pie chart, giving users a visual summary of their monthly financial activity. These charts '
    'are built using the Recharts library and update in real-time as transactions are added.'
)
add_normal_para(
    'Personalized Greeting: The dashboard greets the user by name with a time-appropriate '
    'message (Good morning/evening) and shows the current month\'s financial overview.'
)
add_figure_placeholder('Figure 3.5 Dashboard Overview', 'Screenshot 2026-07-13 205528.png')


# Module 3: Expense Management
add_heading_custom('3.4.3 Expense Management Module', level=3)
add_normal_para(
    'The Expense Management module is the core of the application, allowing users to record, '
    'view, edit, and delete their daily expenses. It provides multiple ways to add expenses '
    'including manual entry, receipt scanning, and voice input.'
)
add_normal_para(
    'Adding Expenses: Users can add expenses by clicking the "+ Add" button which opens a modal '
    'form. The form captures the expense category (Food, Transport, Shopping, Bills, Health, '
    'Entertainment, etc.), description, amount, date, and payment method (Cash, UPI, Card, '
    'Net Banking). Users can also attach receipt images.'
)
add_normal_para(
    'Receipt Scanning: The "Scan" button allows users to upload a receipt or bill image. The '
    'system uses Tesseract.js OCR engine on the backend to extract text from the image and '
    'automatically identify the merchant name, amount, date, and category. This significantly '
    'reduces manual data entry effort.'
)
add_normal_para(
    'Voice Entry: The "Voice" button activates the browser\'s Web Speech API. Users can speak '
    'naturally, for example "I spent 250 rupees on petrol" or "Paid 1500 for groceries", and '
    'the system parses the speech to create an expense entry automatically.'
)
add_normal_para(
    'Export: The "Export" button allows users to download their expense data, making it '
    'useful for record-keeping or tax purposes.'
)
add_normal_para(
    'Search and Filter: A search bar allows users to find specific expenses by description. '
    'A category dropdown filter lets users view expenses from a specific category only.'
)
add_normal_para(
    'Expense Table: All expenses are displayed in a table format showing Category, Description, '
    'Date, Payment Method, Amount, and Actions (Edit/Delete). The table supports pagination for '
    'handling large numbers of entries.'
)
add_figure_placeholder('Figure 3.6 Expenses Page', 'Screenshot 2026-07-13 205600.png')


# Module 4: Income Management
add_heading_custom('3.4.4 Income Management Module', level=3)
add_normal_para(
    'The Income Management module allows users to track all their income sources in one place. '
    'It supports both one-time and recurring income entries, making it suitable for salaried '
    'employees, freelancers, and those with multiple income streams.'
)
add_normal_para(
    'Income Summary: At the top of the page, three summary cards display Total This Month '
    '(total income received in the current month), Transactions (number of income entries), '
    'and Top Source (the highest contributing income source).'
)
add_normal_para(
    'Adding Income: Users click "+ Add Income" to open a form where they can enter the source '
    '(Salary, Freelancing, Investments, Business, Rental, etc.), description, amount, date, '
    'and whether it is a recurring income.'
)
add_normal_para(
    'Recurring Income: For recurring income like monthly salary, users can mark the entry as '
    'recurring so they do not need to manually add it every month.'
)
add_normal_para(
    'Search and Filter: Similar to expenses, the income page provides search by description '
    'and filter by source functionality.'
)
add_normal_para(
    'Income Table: All income entries are displayed in a table with columns for Source, '
    'Description, Date, Recurring status, Amount, and Actions (Edit/Delete).'
)
add_figure_placeholder('Figure 3.7 Income Page', 'Screenshot 2026-07-13 205621.png')

# Module 5: Budget Planning
add_heading_custom('3.4.5 Budget Planning Module', level=3)
add_normal_para(
    'The Budget Planning module enables users to set monthly spending limits for different '
    'expense categories. This helps users control their spending and avoid overspending in '
    'any particular area.'
)
add_normal_para(
    'Creating a Budget: Users click "+ Create Budget" to set up their monthly budget. They '
    'can allocate specific amounts to different categories like Food, Transport, Shopping, '
    'Entertainment, Bills, and others. The total budget gives an overall monthly spending limit.'
)
add_normal_para(
    'Budget Tracking: Once a budget is set, the system automatically tracks actual spending '
    'against the budgeted amount for each category. Progress bars show how much of each '
    'category\'s budget has been used.'
)
add_normal_para(
    'Budget Alerts: When spending in any category approaches or exceeds the set limit, the '
    'notification system alerts the user, helping them stay within their planned budget.'
)
add_normal_para(
    'Monthly Reset: Budgets are created on a monthly basis, and users can view their current '
    'month\'s budget status at any time.'
)
add_figure_placeholder('Figure 3.8 Budget Planning Page', 'Screenshot 2026-07-13 205641.png')


# Module 6: Goal Planning
add_heading_custom('3.4.6 Financial Goals Module', level=3)
add_normal_para(
    'The Financial Goals module helps users set savings targets and track their progress '
    'towards achieving them. Whether saving for a vacation, buying a gadget, or building an '
    'emergency fund, this module keeps users motivated by showing their progress.'
)
add_normal_para(
    'Creating Goals: Users click "+ New Goal" to create a financial goal. They provide a '
    'goal name (e.g., "New Laptop", "Emergency Fund"), target amount, and target date. The '
    'system uses AI to estimate the required monthly savings to achieve the goal on time.'
)
add_normal_para(
    'Contributing to Goals: Users can add contributions towards their goals at any time. '
    'Each contribution is recorded with the date and amount, building a history of savings.'
)
add_normal_para(
    'Goal Status Tabs: Goals are organized into four tabs - Active (currently being saved for), '
    'Achieved (target reached), Paused (temporarily stopped), and All (complete list). This '
    'makes it easy to manage multiple goals simultaneously.'
)
add_normal_para(
    'Progress Tracking: Each goal shows a progress bar indicating how close the user is to '
    'reaching the target amount. The percentage completion and remaining amount are clearly displayed.'
)
add_figure_placeholder('Figure 3.9 Financial Goals Page', 'Screenshot 2026-07-13 205704.png')

# Module 7: Analytics
add_heading_custom('3.4.7 Analytics Module', level=3)
add_normal_para(
    'The Analytics module provides visual representations of the user\'s financial data '
    'through interactive charts. This helps users understand their spending patterns, identify '
    'trends, and make informed financial decisions.'
)
add_normal_para(
    'Monthly Expense Trend: A line chart showing the total expenses for each of the last '
    '12 months. This helps users identify months where they spent more than usual.'
)
add_normal_para(
    'Income vs Expense: A comparison chart showing both income and expenses over the last '
    '12 months on the same graph. Green represents income and red represents expenses, making '
    'it easy to see whether the user is saving or overspending each month.'
)
add_normal_para(
    'Expense by Category: A pie chart or donut chart breaking down expenses by category for '
    'the current month. This shows which categories consume the most money.'
)
add_normal_para(
    'Savings Trend: A line chart showing the monthly savings (income minus expenses) over the '
    'last 12 months. This shows the user\'s saving trajectory over time.'
)
add_normal_para(
    'All charts are built using the Recharts library and are interactive - users can hover '
    'over data points to see exact values. The charts automatically update as new transactions '
    'are added.'
)
add_figure_placeholder('Figure 3.10 Analytics Dashboard', 'Screenshot 2026-07-13 205724.png')


# Module 8: AI Chat
add_heading_custom('3.4.8 AI Chat Module (Gemini AI)', level=3)
add_normal_para(
    'The AI Chat module is one of the most innovative features of FinGenius AI. It integrates '
    'Google Gemini 1.5 Flash as a personal financial advisor that has access to the user\'s '
    'actual financial data and can provide contextual, personalized advice.'
)
add_normal_para(
    'How It Works: When a user sends a message, the backend collects the user\'s financial '
    'summary including monthly expenses by category, income overview, budget status, active '
    'goals progress, and the last five transactions. This context is sent along with the '
    'user\'s question to the Gemini AI model, which generates a relevant response.'
)
add_normal_para(
    'Sample Queries: Users can ask questions like "How much did I spend this month?", "Show '
    'my food expenses", "How can I save more money?", "Can I afford a 70,000 rupee laptop?", '
    '"Summarize my expenses", or "Suggest budget improvements". The AI responds with specific '
    'data-driven answers rather than generic advice.'
)
add_normal_para(
    'Chat History: The module maintains a conversation history in the sidebar. Users can start '
    'new chats or continue previous conversations. Chat messages are stored in the database '
    'for persistence across sessions.'
)
add_normal_para(
    'Suggested Prompts: For new users, the chat interface displays pre-written suggested '
    'prompts to help them understand what questions they can ask.'
)
add_normal_para(
    'Disclaimer: The interface clearly states that "AI responses are based on your financial '
    'data. Not financial advice." This ensures users understand the advisory nature of responses.'
)
add_figure_placeholder('Figure 3.11 AI Chat Interface', 'Screenshot 2026-07-13 205745.png')


# Module 9: Receipt OCR
add_heading_custom('3.4.9 Receipt OCR Module', level=3)
add_normal_para(
    'The Receipt OCR module provides automated expense extraction from bill and receipt images. '
    'Instead of manually typing expense details, users can simply photograph or scan their '
    'receipt and the system will extract the relevant information automatically.'
)
add_normal_para(
    'Technology: The module uses Tesseract.js, an open-source OCR (Optical Character Recognition) '
    'engine that runs on the server. When a user uploads a receipt image, it is processed by '
    'Multer for file handling and then passed to the OCR service for text extraction.'
)
add_normal_para(
    'Extracted Information: The OCR system attempts to identify and extract the merchant name, '
    'total amount, date of purchase, GST details, and automatically suggests an expense '
    'category based on the merchant type.'
)
add_normal_para(
    'User Confirmation: After extraction, the details are pre-filled in the expense form for '
    'the user to review and confirm before saving. Users can correct any errors in the '
    'extracted data before submission.'
)

# Module 10: Voice Entry
add_heading_custom('3.4.10 Voice Entry Module', level=3)
add_normal_para(
    'The Voice Entry module allows users to log expenses by simply speaking, making it a '
    'hands-free and quick method of recording transactions. This is especially useful when '
    'users are on the go and cannot type.'
)
add_normal_para(
    'Technology: The module uses the Web Speech API which is natively supported in modern '
    'browsers like Chrome and Edge. The speech recognition runs entirely in the browser '
    'without sending audio to any external server.'
)
add_normal_para(
    'Natural Language Parsing: Users can speak in natural language, for example "I spent 250 '
    'rupees on petrol" or "Paid 1500 for groceries yesterday". The system parses the spoken '
    'text to identify the amount, category, and description, and creates the expense entry.'
)
add_normal_para(
    'Browser Support: The voice feature works best in Google Chrome and Microsoft Edge. A '
    'microphone icon button on the Expenses page activates the speech recognition.'
)

# Module 11: Notifications
add_heading_custom('3.4.11 Notification System', level=3)
add_normal_para(
    'The Notification System keeps users informed about important events related to their '
    'finances. A bell icon in the top navigation bar shows the notification count and provides '
    'access to the notification panel.'
)
add_normal_para(
    'Types of Notifications: The system generates notifications for budget alerts (when '
    'spending approaches or exceeds the set limit), goal achievements (when a savings goal is '
    'reached), and general reminders.'
)
add_normal_para(
    'Real-time Updates: Notifications are generated automatically by the backend when certain '
    'conditions are met, such as spending exceeding 80% of a category budget.'
)


# Module 12: Settings and Profile
add_heading_custom('3.4.12 Settings and Profile Module', level=3)
add_normal_para(
    'The Settings module allows users to manage their profile information and application '
    'preferences. Users can update their name, email, and other profile details.'
)
add_normal_para(
    'Theme Toggle: The application supports both dark mode and light mode. Users can switch '
    'between themes using the theme toggle in settings. The selected theme preference is '
    'persisted in local storage.'
)
add_normal_para(
    'Profile Management: Users can view and update their profile details through the '
    'settings page. The API supports GET and PUT operations for profile data.'
)

# Module 13: API Layer
add_heading_custom('3.4.13 API Layer', level=3)
add_normal_para(
    'The API layer is the backbone of the application, providing RESTful endpoints for all '
    'frontend operations. It is built with Express.js and follows a modular route-controller '
    'pattern.'
)
add_normal_para(
    'Key API endpoints include:'
)
add_normal_para('- POST /api/auth/register - User registration')
add_normal_para('- POST /api/auth/login - User login')
add_normal_para('- GET /api/auth/me - Get current user details')
add_normal_para('- POST /api/auth/forgot-password - Request password reset')
add_normal_para('- GET /api/dashboard - Complete financial summary')
add_normal_para('- GET/POST /api/expenses - List and create expenses')
add_normal_para('- PUT/DELETE /api/expenses/:id - Update and delete expenses')
add_normal_para('- GET/POST /api/income - List and create income entries')
add_normal_para('- GET/POST /api/budget - List and create budgets')
add_normal_para('- GET /api/budget/current - Current month budget status')
add_normal_para('- GET/POST /api/goals - List and create goals')
add_normal_para('- POST /api/goals/:id/contribute - Add contribution to a goal')
add_normal_para('- POST /api/chat - Send message to AI chatbot')
add_normal_para('- POST /api/ocr/receipt - Process receipt image')
add_normal_para('- GET /api/analytics/* - Various analytics endpoints')
add_normal_para('- GET/PUT /api/profile - Get and update profile')
add_normal_para('- GET /api/notifications - Get user notifications')
add_normal_para(
    'All protected endpoints require a valid JWT token in the Authorization header. The '
    'middleware validates the token and extracts the user ID before processing the request.'
)

# Module 14: Security
add_heading_custom('3.4.14 Security Module', level=3)
add_normal_para(
    'Security is a critical aspect of any financial application. FinGenius AI implements '
    'multiple layers of security to protect user data.'
)
add_normal_para(
    '- JWT Authentication: Tokens expire after 7 days, requiring re-login.'
)
add_normal_para(
    '- Password Hashing: bcryptjs with 12 salt rounds ensures passwords are stored securely.'
)
add_normal_para(
    '- Rate Limiting: API requests are limited to 100 per 15 minutes for general endpoints '
    'and 10 per 15 minutes for authentication endpoints to prevent brute force attacks.'
)
add_normal_para(
    '- CORS: Cross-Origin Resource Sharing is restricted to the client URL only.'
)
add_normal_para(
    '- Helmet.js: Sets various HTTP security headers to prevent common web vulnerabilities.'
)
add_normal_para(
    '- Input Validation: express-validator validates all incoming request data.'
)
add_normal_para(
    '- MongoDB Injection Prevention: Mongoose ODM with schema validation prevents injection attacks.'
)
doc.add_page_break()


# 3.5 Database Design
add_heading_custom('3.5 Database Design', level=2)
add_normal_para(
    'FinGenius AI uses MongoDB, a NoSQL document database, hosted on MongoDB Atlas cloud. '
    'The database consists of seven main collections, each defined by a Mongoose schema. '
    'Below is a description of each collection and its fields.'
)

add_heading_custom('3.5.1 User Collection', level=3)
add_normal_para(
    'The User collection stores all registered user information. Fields include name (String), '
    'email (String, unique), password (String, hashed), createdAt (Date), and resetPasswordToken '
    '(String, for password recovery). Each user document is referenced by other collections '
    'through the userId field.'
)

add_heading_custom('3.5.2 Expense Collection', level=3)
add_normal_para(
    'The Expense collection stores all expense transactions. Fields include userId (ObjectId, '
    'reference to User), category (String - Food, Transport, Shopping, Bills, Health, '
    'Entertainment, Education, Other), description (String), amount (Number), date (Date), '
    'paymentMethod (String - Cash, UPI, Card, Net Banking), receipt (String, file path), and '
    'createdAt (Date).'
)

add_heading_custom('3.5.3 Income Collection', level=3)
add_normal_para(
    'The Income collection stores all income entries. Fields include userId (ObjectId), '
    'source (String - Salary, Freelancing, Investments, Business, Rental, Other), description '
    '(String), amount (Number), date (Date), isRecurring (Boolean), and createdAt (Date).'
)

add_heading_custom('3.5.4 Budget Collection', level=3)
add_normal_para(
    'The Budget collection stores monthly budget allocations. Fields include userId (ObjectId), '
    'month (Number), year (Number), totalBudget (Number), categories (Array of objects with '
    'name and amount), and createdAt (Date). Each user can have one budget per month.'
)

add_heading_custom('3.5.5 Goal Collection', level=3)
add_normal_para(
    'The Goal collection stores financial goals. Fields include userId (ObjectId), name '
    '(String), targetAmount (Number), currentAmount (Number), targetDate (Date), status '
    '(String - active, achieved, paused), contributions (Array of objects with amount and date), '
    'and createdAt (Date).'
)

add_heading_custom('3.5.6 Chat Collection', level=3)
add_normal_para(
    'The Chat collection stores AI chat conversations. Fields include userId (ObjectId), '
    'messages (Array of objects with role - user/assistant, content - String, and timestamp - '
    'Date), title (String, auto-generated from first message), and createdAt (Date).'
)

add_heading_custom('3.5.7 Notification Collection', level=3)
add_normal_para(
    'The Notification collection stores user notifications. Fields include userId (ObjectId), '
    'type (String - budget_alert, goal_achieved, reminder), title (String), message (String), '
    'isRead (Boolean), and createdAt (Date).'
)

add_normal_para('')
add_normal_para(
    'Relationships: All collections are linked to the User collection through the userId field. '
    'This is a reference-based relationship where each document in Expenses, Income, Budget, '
    'Goals, Chat, and Notifications belongs to a specific user. MongoDB\'s flexible schema '
    'allows embedded documents for nested data like budget categories and goal contributions.'
)
add_figure_placeholder('Figure 3.12 Database Schema Diagram')
doc.add_page_break()


# 3.6 Technology Stack
add_heading_custom('3.6 Technology Stack', level=2)
add_normal_para(
    'The following table lists all the technologies, frameworks, and libraries used in the '
    'development of FinGenius AI along with their purpose in the project.'
)

# Create technology stack table
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'S.No'
hdr_cells[1].text = 'Technology'
hdr_cells[2].text = 'Purpose'

# Make header bold
for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

tech_data = [
    ('1', 'React 18', 'Frontend user interface library for building interactive UI components'),
    ('2', 'Vite', 'Build tool and development server for fast frontend development'),
    ('3', 'Tailwind CSS', 'Utility-first CSS framework for responsive glassmorphism design'),
    ('4', 'Framer Motion', 'Animation library for smooth page transitions and UI effects'),
    ('5', 'Recharts', 'Charting library for interactive financial charts and graphs'),
    ('6', 'React Hook Form', 'Form handling library for efficient form validation'),
    ('7', 'React Router v6', 'Client-side routing for single-page application navigation'),
    ('8', 'Axios', 'HTTP client for making API requests from frontend to backend'),
    ('9', 'Node.js', 'JavaScript runtime environment for server-side application'),
    ('10', 'Express.js', 'Web application framework for building RESTful APIs'),
    ('11', 'MongoDB', 'NoSQL document database for storing user and financial data'),
    ('12', 'Mongoose', 'ODM library for MongoDB schema definition and data modeling'),
    ('13', 'JWT (jsonwebtoken)', 'Token-based authentication for secure API access'),
    ('14', 'bcryptjs', 'Password hashing library for secure credential storage'),
    ('15', 'Multer', 'Middleware for handling file uploads (receipt images)'),
    ('16', 'Tesseract.js', 'OCR engine for extracting text from receipt images'),
    ('17', 'Google Gemini AI', 'Large Language Model for AI chatbot and financial advice'),
    ('18', 'Nodemailer', 'Email sending library for password reset functionality'),
    ('19', 'Helmet.js', 'Security middleware for setting HTTP security headers'),
    ('20', 'express-validator', 'Input validation middleware for request data validation'),
    ('21', 'express-rate-limit', 'Rate limiting middleware to prevent abuse'),
    ('22', 'Web Speech API', 'Browser API for voice-based expense entry'),
    ('23', 'MongoDB Atlas', 'Cloud-hosted MongoDB database service'),
]

for sno, tech, purpose in tech_data:
    row_cells = table.add_row().cells
    row_cells[0].text = sno
    row_cells[1].text = tech
    row_cells[2].text = purpose

add_normal_para('')
add_normal_para('Table 3.1 Technology Stack Used in FinGenius AI')
doc.add_page_break()


# ============ CHAPTER 4 - MERITS AND DEMERITS ============
add_heading_custom('CHAPTER 4', level=1)
add_heading_custom('MERITS AND DEMERITS', level=2)

add_heading_custom('4.1 Merits', level=2)
add_normal_para(
    'The following are the advantages of the FinGenius AI application:'
)
add_normal_para(
    '1. AI-Powered Insights: The Gemini AI chatbot provides personalized financial advice '
    'based on actual user data, unlike generic financial tips available elsewhere.'
)
add_normal_para(
    '2. Multiple Input Methods: Users can add expenses through manual entry, receipt scanning '
    '(OCR), or voice input, making the application flexible for different situations.'
)
add_normal_para(
    '3. Completely Free: All features including AI chat, analytics, budgets, and goals are '
    'available without any subscription or payment.'
)
add_normal_para(
    '4. Modern and Intuitive UI: The glassmorphism design with smooth animations and dark/light '
    'mode makes the application visually appealing and easy to use.'
)
add_normal_para(
    '5. Comprehensive Financial Management: Covers expenses, income, budgets, goals, and '
    'analytics in a single platform.'
)
add_normal_para(
    '6. Secure: JWT authentication, password hashing, rate limiting, and input validation '
    'ensure data security.'
)
add_normal_para(
    '7. Responsive Design: Works seamlessly on desktop, tablet, and mobile devices without '
    'needing a separate mobile app.'
)
add_normal_para(
    '8. Privacy-Focused: Does not require access to bank accounts, SMS, or any external '
    'financial data. Users have full control over their data.'
)
add_normal_para(
    '9. Real-time Analytics: Interactive charts update automatically as transactions are added, '
    'providing instant visual feedback.'
)
add_normal_para(
    '10. Cloud-Based: Data is stored in MongoDB Atlas cloud, accessible from anywhere and '
    'protected against local data loss.'
)

add_heading_custom('4.2 Demerits', level=2)
add_normal_para(
    'The following are the current limitations of the system:'
)
add_normal_para(
    '1. No Bank Integration: The application does not automatically fetch transactions from '
    'bank accounts. Users must enter expenses manually or through OCR/voice.'
)
add_normal_para(
    '2. Internet Required: Being a web application with cloud database, internet connectivity '
    'is required for all operations. There is no offline mode.'
)
add_normal_para(
    '3. OCR Accuracy: Receipt scanning depends on image quality. Blurry, damaged, or '
    'handwritten receipts may not be processed accurately.'
)
add_normal_para(
    '4. Voice Recognition Limitations: The Web Speech API voice recognition works best in '
    'English and requires Chrome or Edge browser. It may not work well in noisy environments.'
)
add_normal_para(
    '5. Single User: The application is designed for individual use and does not support '
    'family or shared accounts.'
)
add_normal_para(
    '6. No Mobile App: While responsive, it is a web application and does not have a native '
    'mobile app with push notifications or background syncing.'
)
add_normal_para(
    '7. AI Dependency: The AI chat feature depends on Google Gemini API availability. If the '
    'API is down or the key expires, the chat feature will not work.'
)
doc.add_page_break()


# ============ CHAPTER 5 - COMPARISON WITH EXISTING SYSTEM ============
add_heading_custom('CHAPTER 5', level=1)
add_heading_custom('COMPARISON WITH EXISTING SYSTEM', level=2)
add_normal_para(
    'The table below compares FinGenius AI with traditional manual tracking, spreadsheet-based '
    'tracking, and typical third-party finance applications across various parameters.'
)

comp_table = doc.add_table(rows=1, cols=5)
comp_table.style = 'Table Grid'
comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr = comp_table.rows[0].cells
headers = ['Feature', 'Manual/Notebook', 'Spreadsheet (Excel)', 'Typical Finance Apps', 'FinGenius AI']
for i, h in enumerate(headers):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.bold = True

comp_rows = [
    ('Security', 'None', 'Password-protected file only', 'Varies by app', 'JWT authentication with bcrypt hashing'),
    ('AI Assistance', 'Not available', 'Not available', 'Rare or basic', 'Gemini AI chatbot with personalized advice'),
    ('Expense Entry', 'Manual, slow', 'Manual with formulas', 'Manual entry only', 'Manual, receipt OCR, and voice entry'),
    ('Budget Planning', 'Not available', 'Manual formulas needed', 'Basic limits, often paid', 'Category-wise budgets with alerts'),
    ('Goal Tracking', 'Not available', 'Manual tracking', 'Sometimes available', 'Dedicated goal module with contributions'),
    ('Analytics', 'None', 'Basic charts', 'Standard static charts', 'Interactive Recharts visualizations'),
    ('Notifications', 'Not available', 'Not available', 'Sometimes available', 'Budget alerts and goal notifications'),
    ('Cost', 'Free', 'Free or licensed', 'Free to paid subscription', 'Free for this project'),
    ('Accessibility', 'Physical book only', 'Needs installed software', 'Mobile app required', 'Responsive web app on any device'),
]
for row_data in comp_rows:
    cells = comp_table.add_row().cells
    for i, val in enumerate(row_data):
        cells[i].text = val

add_normal_para('')
add_normal_para('Table 5.1 Comparison of FinGenius AI with Existing Systems')
add_normal_para('')
add_normal_para(
    'As shown in the comparison, manual and spreadsheet-based methods lack any form of '
    'automation, security, or intelligence. While typical finance apps offer better structure '
    'than manual methods, they rarely combine AI-based assistance, receipt scanning, voice '
    'entry, and complete budget-goal-analytics functionality within a single free platform, '
    'as FinGenius AI does.'
)
doc.add_page_break()


# ============ CHAPTER 6 - CONCLUSION AND FUTURE SCOPE ============
add_heading_custom('CHAPTER 6', level=1)
add_heading_custom('CONCLUSION AND FUTURE SCOPE', level=2)

add_heading_custom('6.1 Conclusion', level=2)
add_normal_para(
    'FinGenius AI successfully demonstrates how personal finance management can be made '
    'simpler, smarter, and more accessible by combining modern web technologies with '
    'artificial intelligence. The project implements a complete set of features covering '
    'expense tracking, income management, budget planning, goal setting, analytics, and an '
    'AI-powered chat assistant, all within a single, free, easy-to-use web application.'
)
add_normal_para(
    'The use of the MERN stack (MongoDB, Express.js, React, Node.js) provided a solid and '
    'scalable foundation for building the application, while the integration of Google Gemini '
    'AI added a layer of intelligence that is missing from most existing free finance tools. '
    'The addition of Receipt OCR and Voice Entry further improved the convenience of logging '
    'transactions, addressing one of the biggest reasons people abandon expense tracking '
    'applications - the effort required for manual data entry.'
)
add_normal_para(
    'The project also gave hands-on experience in building a secure, full-stack application '
    'with proper authentication, RESTful API design, database modeling with MongoDB, and '
    'integrating third-party AI services. Overall, FinGenius AI meets its stated objectives '
    'and provides a strong foundation that can be extended further in the future.'
)

add_heading_custom('6.2 Future Scope', level=2)
add_normal_para(
    'While the current version of FinGenius AI covers the core needs of personal finance '
    'tracking, there are several directions in which the project can be extended:'
)
add_normal_para(
    '1. Bank Account Integration: Connecting to bank APIs to automatically import transactions '
    'instead of relying solely on manual entry.'
)
add_normal_para(
    '2. UPI Integration: Adding direct integration with UPI payment systems to automatically '
    'log digital payments.'
)
add_normal_para(
    '3. Improved OCR Accuracy: Enhancing the receipt scanning feature to handle a wider '
    'variety of receipt formats and handwritten bills.'
)
add_normal_para(
    '4. Multi-language Support: Adding support for regional languages in both the UI and the '
    'voice entry feature.'
)
add_normal_para(
    '5. Investment Tracking: Introducing a module to track stocks, mutual funds, and other '
    'investments alongside regular income and expenses.'
)
add_normal_para(
    '6. Predictive Analytics: Building a feature that forecasts future expenses based on past '
    'spending patterns using machine learning.'
)
add_normal_para(
    '7. Mobile Application: Developing a dedicated mobile app with push notifications for a '
    'more native experience.'
)
add_normal_para(
    '8. Family/Shared Accounts: Allowing multiple users to share and manage a common household '
    'budget.'
)
add_normal_para(
    '9. Advanced AI Advisor: Expanding the Gemini AI assistant to provide more advanced, '
    'personalized savings and investment suggestions.'
)
doc.add_page_break()

# ============ REFERENCES ============
add_heading_custom('REFERENCES', level=1)
add_normal_para('')
references = [
    '1. React.js Official Documentation - https://react.dev',
    '2. Node.js Official Documentation - https://nodejs.org/en/docs',
    '3. Express.js Official Documentation - https://expressjs.com',
    '4. MongoDB Official Documentation - https://www.mongodb.com/docs',
    '5. Mongoose Official Documentation - https://mongoosejs.com/docs/guide.html',
    '6. Tailwind CSS Official Documentation - https://tailwindcss.com/docs',
    '7. Recharts Official Documentation - https://recharts.org',
    '8. Google Gemini API Documentation - https://ai.google.dev/gemini-api/docs',
    '9. Tesseract.js Documentation - https://tesseract.projectnaptha.com',
    '10. JSON Web Token (JWT) Introduction - https://jwt.io/introduction',
    '11. Web Speech API - MDN Web Docs - https://developer.mozilla.org/en-US/docs/Web/API/Web_Speech_API',
    '12. Vite Official Documentation - https://vitejs.dev',
    '13. Framer Motion Documentation - https://www.framer.com/motion',
    '14. GitHub Documentation - https://docs.github.com',
]
for ref in references:
    add_normal_para(ref)

# ============ SAVE DOCUMENT ============
output_path = '/projects/sandbox/summa2/FinGenius_AI_Mini_Project_Viva_Report_NEW.docx'
doc.save(output_path)
print(f'Report saved to {output_path}')
